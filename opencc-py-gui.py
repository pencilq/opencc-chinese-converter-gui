#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
OpenCC GUI - Chinese Text Conversion Tool
A GUI application for converting Chinese text using OpenCC
Supports Excel, Word, and Text files with column selection
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import pandas as pd
import os
import threading
from pathlib import Path
import ctypes
try:
    from opencc import OpenCC
except ImportError:
    messagebox.showerror("Import Error", "Please install opencc-python-reimplemented:\npip install opencc-python-reimplemented")
    exit(1)

try:
    from openpyxl import load_workbook
except ImportError:
    messagebox.showerror("Import Error", "Please install openpyxl:\npip install openpyxl")
    exit(1)

try:
    from docx import Document
except ImportError:
    messagebox.showerror("Import Error", "Please install python-docx:\npip install python-docx")
    exit(1)


class OpenCCGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("OpenCC 中文转换器")
        self.root.geometry("1200x700")
        
        # Variables
        self.input_file_path = tk.StringVar()
        self.output_file_path = tk.StringVar()
        # New multi-layer conversion settings
        self.source_type = tk.StringVar(value="简体")
        self.target_type = tk.StringVar(value="繁体")
        self.variant_standard = tk.StringVar(value="台湾标准")
        self.convert_phrases = tk.BooleanVar(value=True)
        self.file_type = tk.StringVar()
        self.direct_text_input = tk.StringVar()
        
        # Data storage
        self.file_data = None
        self.selected_columns = []
        self.column_vars = {}
        
        # OpenCC converter
        self.converter = None
        
        # Auto-preview tracking
        self.last_conversion_settings = ""
        self.last_input_file = ""
        self.last_direct_text = ""
        
        self.setup_ui()
        self.update_converter()
        
        # Set up auto-preview callbacks
        self.input_file_path.trace('w', self.on_input_change)
        self.source_type.trace('w', self.on_conversion_change)
        self.target_type.trace('w', self.on_conversion_change)
        self.variant_standard.trace('w', self.on_conversion_change)
        self.convert_phrases.trace('w', self.on_conversion_change)
        self.direct_text_input.trace('w', self.on_direct_text_change)
    
    def setup_ui(self):
        """Set up the user interface"""
        # Set consistent font for entire application
        app_font = ("Cascadia Code", 9, "normal")
        try:
            self.root.option_add("*Font", app_font)
            # Also set font for specific widget types
            self.root.option_add("*TLabel*Font", app_font)
            self.root.option_add("*TButton*Font", app_font)
            self.root.option_add("*TCheckbutton*Font", app_font)
            self.root.option_add("*TCombobox*Font", app_font)
        except:
            fallback_font = ("黑体", 9, "normal")
            self.root.option_add("*Font", fallback_font)
            self.root.option_add("*TLabel*Font", fallback_font)
            self.root.option_add("*TButton*Font", fallback_font)
            self.root.option_add("*TCheckbutton*Font", fallback_font)
            self.root.option_add("*TCombobox*Font", fallback_font)
        
        # Main frame
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky="nsew")
        
        # Configure grid weights for two-column layout with separator
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)  # Left column (controls)
        main_frame.columnconfigure(1, weight=0)  # Separator (fixed width)
        main_frame.columnconfigure(2, weight=1)  # Right column (preview)
        main_frame.rowconfigure(0, weight=1)    # Allow vertical expansion
        
        # Create left and right containers
        left_frame = ttk.Frame(main_frame, padding=(0, 0, 15, 0))  # Right padding for separation
        left_frame.grid(row=0, column=0, sticky="nsew")
        left_frame.columnconfigure(0, weight=1)
        
        # Visual separator
        separator = ttk.Separator(main_frame, orient="vertical")
        separator.grid(row=0, column=1, sticky="ns", padx=5)
        
        right_frame = ttk.Frame(main_frame, padding=(15, 0, 0, 0))  # Left padding for separation
        right_frame.grid(row=0, column=2, sticky="nsew")
        right_frame.columnconfigure(0, weight=1)
        right_frame.rowconfigure(1, weight=1)  # Preview area should expand
        
        # Title in left frame
        title_label = ttk.Label(left_frame, text="OpenCC 中文转换器", 
                               font=("Cascadia Code", 14, "bold"))
        title_label.grid(row=0, column=0, pady=(0, 20), sticky="w")
        
        # File selection section in left frame
        file_frame = ttk.LabelFrame(left_frame, text="文件选择", padding="10")
        file_frame.grid(row=1, column=0, sticky="ew", pady=(0, 10))
        file_frame.columnconfigure(1, weight=1)
        
        # Input file
        ttk.Label(file_frame, text="输入文件:").grid(row=0, column=0, sticky=tk.W, padx=(0, 5))
        ttk.Entry(file_frame, textvariable=self.input_file_path, state="readonly").grid(
            row=0, column=1, sticky="ew", padx=(0, 5))
        ttk.Button(file_frame, text="浏览", command=self.browse_input_file).grid(row=0, column=2)
        
        # Output file
        ttk.Label(file_frame, text="输出文件:").grid(row=1, column=0, sticky=tk.W, padx=(0, 5), pady=(5, 0))
        ttk.Entry(file_frame, textvariable=self.output_file_path).grid(
            row=1, column=1, sticky="ew", padx=(0, 5), pady=(5, 0))
        ttk.Button(file_frame, text="浏览", command=self.browse_output_file).grid(row=1, column=2, pady=(5, 0))
        
        # Direct text input section in left frame
        text_input_frame = ttk.LabelFrame(left_frame, text="直接文本输入（可替代文件输入）", padding="10")
        text_input_frame.grid(row=2, column=0, sticky="ew", pady=(0, 10))
        text_input_frame.columnconfigure(0, weight=1)
        
        self.direct_text_entry = scrolledtext.ScrolledText(text_input_frame, height=4, wrap=tk.WORD, 
                                                           font=("Cascadia Code", 9, "normal"))
        self.direct_text_entry.grid(row=0, column=0, sticky="ew")
        self.direct_text_entry.bind('<KeyRelease>', self.on_direct_text_change_event)
        
        # Conversion mode section with multi-layer settings in left frame
        mode_frame = ttk.LabelFrame(left_frame, text="转换设置", padding="10")
        mode_frame.grid(row=3, column=0, sticky="ew", pady=(0, 10))
        mode_frame.columnconfigure(1, weight=1)
        mode_frame.columnconfigure(3, weight=1)
        
        # First row: Source and Target
        ttk.Label(mode_frame, text="原文:").grid(row=0, column=0, sticky=tk.W, padx=(0, 5))
        source_combo = ttk.Combobox(mode_frame, textvariable=self.source_type, 
                                   values=["简体", "繁体"], state="readonly", width=10)
        source_combo.grid(row=0, column=1, sticky=tk.W, padx=(0, 20))
        source_combo.bind('<<ComboboxSelected>>', self.on_conversion_settings_change)
        
        ttk.Label(mode_frame, text="目标:").grid(row=0, column=2, sticky=tk.W, padx=(0, 5))
        target_combo = ttk.Combobox(mode_frame, textvariable=self.target_type, 
                                   values=["简体", "繁体"], state="readonly", width=10)
        target_combo.grid(row=0, column=3, sticky=tk.W)
        target_combo.bind('<<ComboboxSelected>>', self.on_conversion_settings_change)
        
        # Second row: Variant and Phrases
        ttk.Label(mode_frame, text="字形:").grid(row=1, column=0, sticky=tk.W, padx=(0, 5), pady=(10, 0))
        variant_combo = ttk.Combobox(mode_frame, textvariable=self.variant_standard, 
                                    values=["OpenCC 标准", "香港标准", "台湾标准"], 
                                    state="readonly", width=15)
        variant_combo.grid(row=1, column=1, sticky=tk.W, padx=(0, 20), pady=(10, 0))
        variant_combo.bind('<<ComboboxSelected>>', self.on_conversion_settings_change)
        
        phrase_check = ttk.Checkbutton(mode_frame, text="当地词汇", variable=self.convert_phrases,
                                      command=self.on_conversion_settings_change)
        phrase_check.grid(row=1, column=2, columnspan=2, sticky=tk.W, pady=(10, 0))
        
        # Column selection section (for Excel files) - Multi-column support in left frame
        self.column_frame = ttk.LabelFrame(left_frame, text="列选择（Excel文件）", padding="10")
        self.column_frame.grid(row=4, column=0, sticky="ew", pady=(0, 10))
        self.column_frame.columnconfigure(0, weight=1)
        
        # Control buttons row
        control_frame = ttk.Frame(self.column_frame)
        control_frame.grid(row=0, column=0, sticky="ew", pady=(0, 5))
        
        ttk.Button(control_frame, text="全选", command=self.select_all_columns).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(control_frame, text="取消全选", command=self.deselect_all_columns).pack(side=tk.LEFT, padx=(0, 10))
        
        self.selected_count_label = ttk.Label(control_frame, text="已选择: 0 列")
        self.selected_count_label.pack(side=tk.LEFT)
        
        # Scrollable checkboxes area
        checkbox_container = ttk.Frame(self.column_frame)
        checkbox_container.grid(row=1, column=0, sticky="nsew", pady=(5, 0))
        checkbox_container.columnconfigure(0, weight=1)
        checkbox_container.rowconfigure(0, weight=1)
        
        self.column_canvas = tk.Canvas(checkbox_container, height=120)
        self.column_scrollbar = ttk.Scrollbar(checkbox_container, orient="vertical", command=self.column_canvas.yview)
        self.column_scrollable_frame = ttk.Frame(self.column_canvas)
        
        self.column_scrollable_frame.bind(
            "<Configure>",
            lambda e: self.column_canvas.configure(scrollregion=self.column_canvas.bbox("all"))
        )
        
        self.column_canvas.create_window((0, 0), window=self.column_scrollable_frame, anchor="nw")
        self.column_canvas.configure(yscrollcommand=self.column_scrollbar.set)
        
        self.column_canvas.grid(row=0, column=0, sticky="nsew")
        self.column_scrollbar.grid(row=0, column=1, sticky="ns")
        
        # Configure grid weights for proper resizing
        self.column_frame.rowconfigure(1, weight=1)
        
        # Buttons section in left frame
        button_frame = ttk.Frame(left_frame)
        button_frame.grid(row=5, column=0, pady=(10, 0), sticky="ew")
        
        ttk.Button(button_frame, text="转换文件", command=self.convert_file).grid(
            row=0, column=0, padx=(0, 10), sticky="w")
        ttk.Button(button_frame, text="清空", command=self.clear_all).grid(row=0, column=1, sticky="w")
        
        # === RIGHT COLUMN: Preview and Progress ===
        
        # Preview section title in right frame
        preview_title = ttk.Label(right_frame, text="预览与复制区域", 
                                 font=("Cascadia Code", 12, "bold"))
        preview_title.grid(row=0, column=0, pady=(0, 10), sticky="w")
        
        # Preview section in right frame
        preview_frame = ttk.LabelFrame(right_frame, text="", padding="10")
        preview_frame.grid(row=1, column=0, sticky="nsew", pady=(0, 10))
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(0, weight=1)
        
        # Set font with fallback
        try:
            preview_font = ("Cascadia Code", 9, "normal")
        except:
            preview_font = ("黑体", 9, "normal")
        
        self.preview_text = scrolledtext.ScrolledText(preview_frame, height=15, wrap=tk.WORD, font=preview_font)
        self.preview_text.grid(row=0, column=0, sticky="nsew")
        
        # Copy button for direct text conversion in right frame
        copy_button_frame = ttk.Frame(preview_frame)
        copy_button_frame.grid(row=1, column=0, sticky="ew", pady=(5, 0))
        ttk.Button(copy_button_frame, text="复制", command=self.copy_converted_text).pack(side=tk.LEFT)
        
        # Progress section in right frame
        progress_frame = ttk.Frame(right_frame)
        progress_frame.grid(row=2, column=0, sticky="ew", pady=(0, 10))
        progress_frame.columnconfigure(0, weight=1)
        
        self.progress_var = tk.StringVar(value="就绪")
        self.progress_label = ttk.Label(progress_frame, textvariable=self.progress_var)
        self.progress_label.grid(row=0, column=0, sticky=tk.W)
        
        self.progress_bar = ttk.Progressbar(progress_frame, mode="determinate")
        self.progress_bar.grid(row=1, column=0, sticky="ew", pady=(5, 0))
    
    def get_conversion_mode(self):
        """Generate OpenCC conversion mode based on current settings"""
        source = self.source_type.get()
        target = self.target_type.get()
        variant = self.variant_standard.get()
        phrases = self.convert_phrases.get()
        
        # Basic conversion mapping
        if source == "简体" and target == "繁体":
            if variant == "台湾标准":
                return "s2twp" if phrases else "s2tw"
            elif variant == "香港标准":
                return "s2hk"
            else:  # OpenCC 标准
                return "s2t"
        elif source == "繁体" and target == "简体":
            if variant == "台湾标准":
                return "tw2sp" if phrases else "tw2s"
            elif variant == "香港标准":
                return "hk2s"
            else:  # OpenCC 标准
                return "t2s"
        elif source == "繁体" and target == "繁体":
            if variant == "台湾标准":
                return "t2tw"
            elif variant == "香港标准":
                return "t2hk"
            else:
                return "t2s"  # Fallback
        else:
            # 简体到简体，直接返回原文
            return None
    
    def get_conversion_mode_for_filename(self):
        """Get conversion mode string for filename"""
        source = self.source_type.get()
        target = self.target_type.get()
        variant = self.variant_standard.get()
        phrases = self.convert_phrases.get()
        
        # Generate descriptive filename suffix
        if source == target == "简体":
            return "简体"
        elif source == "简体" and target == "繁体":
            if variant == "台湾标准":
                return "s2twp" if phrases else "s2tw"
            elif variant == "香港标准":
                return "s2hk"
            else:
                return "s2t"
        elif source == "繁体" and target == "简体":
            if variant == "台湾标准":
                return "tw2sp" if phrases else "tw2s"
            elif variant == "香港标准":
                return "hk2s"
            else:
                return "t2s"
        elif source == "繁体" and target == "繁体":
            if variant == "台湾标准":
                return "t2tw"
            elif variant == "香港标准":
                return "t2hk"
            else:
                return "t2t"
        else:
            return "convert"
    
    def update_converter(self):
        """Update the OpenCC converter based on current settings"""
        try:
            conversion_mode = self.get_conversion_mode()
            if conversion_mode:
                self.converter = OpenCC(conversion_mode)
                self.progress_var.set(f"转换器就绪: {conversion_mode}")
            else:
                self.converter = None
                self.progress_var.set("无需转换（相同格式）")
        except Exception as e:
            messagebox.showerror("转换器错误", f"初始化转换器失败: {str(e)}")
    
    def on_conversion_settings_change(self, *args):
        """Handle conversion settings change"""
        self.update_converter()
        self.update_output_filename()  # Update filename when settings change
        self.auto_preview()
    
    def update_output_filename(self):
        """Update output filename based on current conversion settings"""
        input_path = self.input_file_path.get()
        if input_path:
            input_path_obj = Path(input_path)
            conversion_mode = self.get_conversion_mode_for_filename()
            output_name = f"{input_path_obj.stem}_{conversion_mode}{input_path_obj.suffix}"
            output_path = input_path_obj.parent / output_name
            self.output_file_path.set(str(output_path))
    
    def on_conversion_mode_change(self):
        """Handle conversion mode change (legacy compatibility)"""
        self.update_converter()
        self.auto_preview()
    
    def on_input_change(self, *args):
        """Handle input file change"""
        self.auto_preview()
    
    def on_conversion_change(self, *args):
        """Handle conversion mode change from trace"""
        current_settings = f"{self.source_type.get()}-{self.target_type.get()}-{self.variant_standard.get()}-{self.convert_phrases.get()}"
        if current_settings != self.last_conversion_settings:
            self.last_conversion_settings = current_settings
            self.update_converter()
            self.update_output_filename()  # Update filename when settings change
            self.auto_preview()
    
    def on_direct_text_change(self, *args):
        """Handle direct text input change from trace"""
        self.auto_preview()
    
    def on_direct_text_change_event(self, event):
        """Handle direct text input change from event"""
        self.auto_preview()
    
    def on_column_selection_change(self, *args):
        """Handle column selection change"""
        self.update_selection_count()
        self.auto_preview()
    
    def copy_converted_text(self):
        """Copy converted text to clipboard"""
        try:
            converted_text = self.preview_text.get(1.0, tk.END).strip()
            if converted_text:
                self.root.clipboard_clear()
                self.root.clipboard_append(converted_text)
                self.root.update()  # Ensure clipboard is updated
                messagebox.showinfo("成功", "转换结果已复制到剪贴板！")
            else:
                messagebox.showwarning("警告", "没有转换结果可复制。")
        except Exception as e:
            messagebox.showerror("错误", f"复制文本失败: {str(e)}")
    
    def browse_input_file(self):
        """Browse for input file"""
        file_path = filedialog.askopenfilename(
            title="Select input file",
            filetypes=[
                ("All supported", "*.xlsx;*.xls;*.docx;*.txt"),
                ("Excel files", "*.xlsx;*.xls"),
                ("Word documents", "*.docx"),
                ("Text files", "*.txt"),
                ("All files", "*.*")
            ]
        )
        
        if file_path:
            self.input_file_path.set(file_path)
            self.detect_file_type(file_path)
            self.load_file_data(file_path)
            
            # Auto-suggest output file name with conversion mode
            input_path = Path(file_path)
            conversion_mode = self.get_conversion_mode_for_filename()
            output_name = f"{input_path.stem}_{conversion_mode}{input_path.suffix}"
            output_path = input_path.parent / output_name
            self.output_file_path.set(str(output_path))
    
    def browse_output_file(self):
        """Browse for output file"""
        input_path = self.input_file_path.get()
        if not input_path:
            messagebox.showwarning("警告", "请先选择输入文件。")
            return
        
        # Determine file type for save dialog
        file_ext = Path(input_path).suffix.lower()
        if file_ext in ['.xlsx', '.xls']:
            filetypes = [("Excel 文件", "*.xlsx"), ("所有文件", "*.*")]
            default_ext = ".xlsx"
        elif file_ext == '.docx':
            filetypes = [("Word 文档", "*.docx"), ("所有文件", "*.*")]
            default_ext = ".docx"
        else:
            filetypes = [("文本文件", "*.txt"), ("所有文件", "*.*")]
            default_ext = ".txt"
        
        file_path = filedialog.asksaveasfilename(
            title="保存转换后的文件",
            filetypes=filetypes,
            defaultextension=default_ext
        )
        
        if file_path:
            self.output_file_path.set(file_path)
    
    def detect_file_type(self, file_path):
        """Detect file type based on extension"""
        ext = Path(file_path).suffix.lower()
        if ext in ['.xlsx', '.xls']:
            self.file_type.set('excel')
        elif ext == '.docx':
            self.file_type.set('word')
        elif ext == '.txt':
            self.file_type.set('text')
        else:
            self.file_type.set('unknown')
    
    def load_file_data(self, file_path):
        """Load file data and update column selection if needed"""
        try:
            self.progress_var.set("Loading file...")
            
            if self.file_type.get() == 'excel':
                # Load Excel file
                self.file_data = pd.read_excel(file_path)
                self.update_column_selection()
            elif self.file_type.get() == 'word':
                # Load Word file
                doc = Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                self.file_data = text
                self.clear_column_selection()
            elif self.file_type.get() == 'text':
                # Load text file
                with open(file_path, 'r', encoding='utf-8') as f:
                    self.file_data = f.read()
                self.clear_column_selection()
            
            self.progress_var.set("File loaded successfully")
        except Exception as e:
            messagebox.showerror("File Error", f"Failed to load file: {str(e)}")
            self.progress_var.set("Error loading file")
    
    def update_column_selection(self):
        """Update column selection checkboxes for Excel files"""
        # Clear existing checkboxes
        for widget in self.column_scrollable_frame.winfo_children():
            widget.destroy()
        
        self.column_vars.clear()
        
        if self.file_data is not None and hasattr(self.file_data, 'columns'):
            columns = self.file_data.columns.tolist()
            for i, col in enumerate(columns):
                var = tk.BooleanVar()
                var.trace('w', self.on_column_selection_change)
                self.column_vars[col] = var
                
                # Create frame for each checkbox with sample data
                checkbox_frame = ttk.Frame(self.column_scrollable_frame)
                checkbox_frame.grid(row=i, column=0, sticky="ew", padx=5, pady=2)
                checkbox_frame.columnconfigure(1, weight=1)
                
                # Checkbox
                cb = ttk.Checkbutton(checkbox_frame, variable=var)
                cb.grid(row=0, column=0, sticky="w")
                
                # Column name and sample data
                sample_text = ""
                if len(self.file_data) > 0:
                    try:
                        sample_val = str(self.file_data[col].iloc[0])
                        sample_text = sample_val[:20] + "..." if len(sample_val) > 20 else sample_val
                    except:
                        sample_text = "N/A"
                
                col_label = ttk.Label(checkbox_frame, 
                                     text=f"{col} (示例: {sample_text})")
                col_label.grid(row=0, column=1, sticky="w", padx=(5, 0))
                
                # Make the checkbox frame expand
                self.column_scrollable_frame.rowconfigure(i, weight=0)
            
            self.column_scrollable_frame.columnconfigure(0, weight=1)
            self.update_selection_count()
        else:
            # Show message when no Excel data
            no_data_label = ttk.Label(self.column_scrollable_frame, text="没有可用列")
            no_data_label.grid(row=0, column=0, pady=20)
    
    def clear_column_selection(self):
        """Clear column selection for non-Excel files"""
        # Clear existing checkboxes
        for widget in self.column_scrollable_frame.winfo_children():
            widget.destroy()
        self.column_vars.clear()
        
        # Show message for non-Excel files
        no_data_label = ttk.Label(self.column_scrollable_frame, text="不适用于此文件类型")
        no_data_label.grid(row=0, column=0, pady=20)
        self.update_selection_count()
    
    def get_selected_columns(self):
        """Get list of selected columns"""
        return [col for col, var in self.column_vars.items() if var.get()]
    
    def select_all_columns(self):
        """Select all available columns"""
        for var in self.column_vars.values():
            var.set(True)
        self.update_selection_count()
        self.auto_preview()
    
    def deselect_all_columns(self):
        """Deselect all columns"""
        for var in self.column_vars.values():
            var.set(False)
        self.update_selection_count()
        self.auto_preview()
    
    def update_selection_count(self):
        """Update the selected columns count display"""
        selected_count = len(self.get_selected_columns())
        total_count = len(self.column_vars)
        if hasattr(self, 'selected_count_label'):
            self.selected_count_label.config(text=f"已选择: {selected_count}/{total_count} 列")
    
    def auto_preview(self):
        """Automatically preview conversion based on current inputs"""
        if not self.converter:
            self.preview_text.delete(1.0, tk.END)
            self.preview_text.insert(tk.END, "转换器未初始化或无需转换。")
            return
        
        try:
            self.preview_text.delete(1.0, tk.END)
            
            # Check for direct text input first
            direct_text = self.direct_text_entry.get(1.0, tk.END).strip()
            if direct_text:
                converted_text = self.converter.convert(direct_text)
                # Only show converted result
                self.preview_text.insert(tk.END, converted_text)
                return
            
            # Check for file input
            if not self.input_file_path.get() or self.file_data is None:
                self.preview_text.insert(tk.END, "没有提供输入。请选择文件或在上方直接输入文本。")
                return
            
            # Check if file_data is empty DataFrame
            if hasattr(self.file_data, 'empty') and self.file_data.empty:
                self.preview_text.insert(tk.END, "文件为空或没有数据。")
                return
            
            if self.file_type.get() == 'excel':
                if not hasattr(self.file_data, 'columns'):
                    self.preview_text.insert(tk.END, "无效的 Excel 数据。")
                    return
                    
                selected_cols = self.get_selected_columns()
                if not selected_cols:
                    available_cols = ", ".join(self.file_data.columns.tolist())
                    self.preview_text.insert(tk.END, f"请选择至少一列进行转换。\n\n可用列: {available_cols}")
                    return
                
                preview_data = self.file_data.head(5).copy()
                for col in selected_cols:
                    if col in preview_data.columns:
                        preview_data[col] = preview_data[col].astype(str).apply(
                            lambda x: self.converter.convert(x) if pd.notna(x) and str(x) != 'nan' and str(x).strip() else x
                        )
                
                # Display results for multiple columns
                if len(selected_cols) == 1:
                    self.preview_text.insert(tk.END, f"列 '{selected_cols[0]}' 的转换结果（前5行）：\n")
                    self.preview_text.insert(tk.END, "=" * 40 + "\n")
                    # Only show the converted column values
                    for idx, value in preview_data[selected_cols[0]].items():
                        self.preview_text.insert(tk.END, f"{str(value)}\n")
                else:
                    self.preview_text.insert(tk.END, f"多列转换结果（{len(selected_cols)}列，前5行）：\n")
                    self.preview_text.insert(tk.END, "=" * 50 + "\n")
                    
                    # Show column headers
                    headers = " | ".join([f"{col[:15]}..." if len(col) > 15 else col for col in selected_cols])
                    self.preview_text.insert(tk.END, f"{headers}\n")
                    self.preview_text.insert(tk.END, "-" * 50 + "\n")
                    
                    # Show data rows
                    for idx in range(min(5, len(preview_data))):
                        row_values = []
                        for col in selected_cols:
                            value = str(preview_data[col].iloc[idx])
                            # Truncate long values for display
                            if len(value) > 15:
                                value = value[:12] + "..."
                            row_values.append(value)
                        
                        row_text = " | ".join(row_values)
                        self.preview_text.insert(tk.END, f"{row_text}\n")
                
            else:  # Word or Text files
                if isinstance(self.file_data, str):
                    sample_text = self.file_data[:500] if len(self.file_data) > 500 else self.file_data
                    converted_text = self.converter.convert(sample_text)
                    
                    # Only show converted result
                    self.preview_text.insert(tk.END, "文件转换结果（前500字符）：\n")
                    self.preview_text.insert(tk.END, "=" * 30 + "\n")
                    self.preview_text.insert(tk.END, converted_text)
                else:
                    self.preview_text.insert(tk.END, "无效的文本数据。")
        
        except Exception as e:
            self.preview_text.insert(tk.END, f"预览错误: {str(e)}")
    
    def convert_file(self):
        """Convert the entire file or direct text input"""
        # Check for direct text input first
        direct_text = self.direct_text_entry.get(1.0, tk.END).strip()
        if direct_text:
            try:
                if not self.converter:
                    messagebox.showerror("错误", "转换器未初始化。")
                    return
                
                converted_text = self.converter.convert(direct_text)
                self.preview_text.delete(1.0, tk.END)
                # Only show converted result
                self.preview_text.insert(tk.END, converted_text)
                self.progress_var.set("直接文本转换完成！")
                messagebox.showinfo("成功", "文本转换成功！您可以从预览区域复制结果。")
                return
            except Exception as e:
                messagebox.showerror("转换错误", f"文本转换失败: {str(e)}")
                return
        
        # Otherwise proceed with file conversion
        if not self.input_file_path.get():
            messagebox.showwarning("警告", "请选择输入文件或直接输入文本。")
            return
        
        if not self.output_file_path.get():
            messagebox.showwarning("警告", "请指定输出文件。")
            return
        
        if not self.converter:
            messagebox.showerror("错误", "转换器未初始化。")
            return
        
        # Run conversion in separate thread to avoid freezing GUI
        thread = threading.Thread(target=self._convert_file_worker)
        thread.daemon = True
        thread.start()
    
    def _convert_file_worker(self):
        """Worker function for file conversion"""
        try:
            # Update UI in main thread
            self.root.after(0, lambda: self.progress_bar.configure(mode="determinate"))
            self.root.after(0, lambda: self.progress_bar.configure(value=0))
            self.root.after(0, lambda: self.progress_var.set("正在转换文件..."))
            
            if self.file_type.get() == 'excel':
                if not hasattr(self.file_data, 'columns'):
                    self.root.after(0, lambda: messagebox.showerror("错误", "无效的 Excel 数据。"))
                    return
                    
                selected_cols = self.get_selected_columns()
                if not selected_cols:
                    self.root.after(0, lambda: messagebox.showwarning("警告", "请选择至少一列进行转换。"))
                    return
                
                # Convert Excel file - multiple columns
                converted_data = self.file_data.copy()
                total_rows = len(converted_data)
                total_operations = total_rows * len(selected_cols)
                current_operation = 0
                
                for col_idx, col in enumerate(selected_cols):
                    if col in converted_data.columns:
                        for row_idx in range(len(converted_data)):
                            value = converted_data.iloc[row_idx, converted_data.columns.get_loc(col)]
                            if pd.notna(value) and str(value) != 'nan' and str(value).strip():
                                converted_data.iloc[row_idx, converted_data.columns.get_loc(col)] = self.converter.convert(str(value))
                            
                            current_operation += 1
                            # Update progress
                            progress_percentage = int(current_operation / total_operations * 100)
                            progress_msg = f"正在转换列 '{col}' ({col_idx+1}/{len(selected_cols)}): {row_idx+1}/{total_rows} ({progress_percentage}%)"
                            self.root.after(0, lambda msg=progress_msg: self.progress_var.set(msg))
                            self.root.after(0, lambda pct=progress_percentage: self.progress_bar.configure(value=pct))
                
                # Save Excel file
                converted_data.to_excel(self.output_file_path.get(), index=False)
                
            elif self.file_type.get() == 'word':
                # Convert Word file
                doc = Document(self.input_file_path.get())
                
                # Count total elements for progress
                total_elements = len(doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        total_elements += len(row.cells)
                
                processed = 0
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        original_text = paragraph.text
                        converted_text = self.converter.convert(original_text)
                        paragraph.text = converted_text
                    
                    processed += 1
                    progress_percentage = int(processed / total_elements * 100)
                    self.root.after(0, lambda pct=progress_percentage: self.progress_bar.configure(value=pct))
                    self.root.after(0, lambda p=processed, t=total_elements: self.progress_var.set(f"正在转换段落: {p}/{t}"))
                
                # Handle tables in Word document
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                original_text = cell.text
                                converted_text = self.converter.convert(original_text)
                                cell.text = converted_text
                            
                            processed += 1
                            progress_percentage = int(processed / total_elements * 100)
                            self.root.after(0, lambda pct=progress_percentage: self.progress_bar.configure(value=pct))
                            self.root.after(0, lambda p=processed, t=total_elements: self.progress_var.set(f"正在转换表格: {p}/{t}"))
                
                doc.save(self.output_file_path.get())
                
            elif self.file_type.get() == 'text':
                # Convert text file
                if isinstance(self.file_data, str):
                    self.root.after(0, lambda: self.progress_var.set("正在转换文本文件..."))
                    self.root.after(0, lambda: self.progress_bar.configure(value=50))
                    
                    converted_text = self.converter.convert(self.file_data)
                    
                    self.root.after(0, lambda: self.progress_bar.configure(value=80))
                    self.root.after(0, lambda: self.progress_var.set("正在保存转换后的文件..."))
                    
                    with open(self.output_file_path.get(), 'w', encoding='utf-8') as f:
                        f.write(converted_text)
                    
                    self.root.after(0, lambda: self.progress_bar.configure(value=100))
                else:
                    self.root.after(0, lambda: messagebox.showerror("错误", "无效的文本数据。"))
                    return
            
            # Success
            self.root.after(0, lambda: self.progress_bar.configure(value=100))
            self.root.after(0, lambda: self.progress_var.set("转换完成！"))
            self.root.after(0, lambda: messagebox.showinfo("成功", f"文件转换成功！\n已保存至: {self.output_file_path.get()}"))
            
        except Exception as e:
            self.root.after(0, lambda: self.progress_bar.configure(value=0))
            self.root.after(0, lambda: self.progress_var.set("转换失败"))
            self.root.after(0, lambda: messagebox.showerror("转换错误", f"文件转换失败: {str(e)}"))
    
    def clear_all(self):
        """Clear all inputs and reset the interface"""
        self.input_file_path.set("")
        self.output_file_path.set("")
        self.file_data = None
        self.direct_text_entry.delete(1.0, tk.END)
        self.preview_text.delete(1.0, tk.END)
        self.clear_column_selection()
        # Reset column selection count display
        if hasattr(self, 'selected_count_label'):
            self.selected_count_label.config(text="已选择: 0 列")
        self.progress_var.set("就绪")
        self.progress_bar.configure(value=0)
        
        # Reset conversion settings to defaults
        self.source_type.set("简体")
        self.target_type.set("繁体")
        self.variant_standard.set("台湾标准")
        self.convert_phrases.set(True)
        
        # Reset tracking variables
        self.last_conversion_settings = ""
        self.last_input_file = ""
        self.last_direct_text = ""
        
        # Update converter
        self.update_converter()


def main():
    """Main function to run the application"""
    # Hide console window on Windows (similar to nps-auto.py)
    if os.name == 'nt':  # Windows
        try:
            ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 0)
        except Exception as e:
            print(f"Warning: Failed to hide console window: {e}")
    
    root = tk.Tk()
    app = OpenCCGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()